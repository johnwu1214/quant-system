#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成量化交易系统试运行报告并发送到飞书
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# 创建Word文档
doc = Document()

# 标题
title = doc.add_heading('🦞 量化交易系统 v3.0 试运行分析报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 生成时间
p = doc.add_paragraph()
p.add_run('📅 生成时间: 2026-03-03').bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# 一、系统概述
doc.add_heading('一、系统概述', level=1)
doc.add_paragraph('量化交易系统 v3.0 (AI增强版) 试运行分析报告，包含AI因子挖掘、AI策略生成、财报分析、风控系统等核心功能。')

# 二、数据接口状态
doc.add_heading('二、数据接口状态', level=1)

table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'

# 表头
header = table.rows[0].cells
header[0].text = '数据源'
header[1].text = '接口'
header[2].text = '状态'

# 数据
data = [
    ('MomaAPI', '实时行情', '✅ 正常'),
    ('MomaAPI', '日线数据', '✅ 正常'),
    ('AkShare', '指数数据', '✅ 正常'),
    ('AkShare', '个股K线', '⚠️ 被限制'),
]

for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    for j, cell_data in enumerate(row_data):
        row[j].text = cell_data

# 三、持仓分析
doc.add_heading('三、持仓分析', level=1)

table2 = doc.add_table(rows=5, cols=5)
table2.style = 'Table Grid'

header2 = table2.rows[0].cells
header2[0].text = '股票'
header2[1].text = '代码'
header2[2].text = '当前价'
header2[3].text = '成本价'
header2[4].text = '盈亏'

positions = [
    ('克莱机电', '603960', '¥25.93', '¥17.41', '+12,780元 (+49.0%)'),
    ('中国海油', '600938', '¥43.41', '¥18.51', '+7,470元 (+134.5%)'),
    ('达仁堂', '600329', '¥41.73', '¥41.37', '+180元 (+0.4%)'),
    ('葵花药业', '002737', '¥13.67', '¥14.82', '-2,875元 (-7.7%)'),
]

for i, pos in enumerate(positions):
    row = table2.rows[i+1].cells
    for j, cell_data in enumerate(pos):
        row[j].text = cell_data

# 四、资产汇总
doc.add_heading('四、资产汇总', level=1)

summary = [
    ('总市值', '¥106,983'),
    ('总成本', '¥89,403'),
    ('总盈亏', '+17,555元 (+19.6%)'),
    ('持仓', '4只'),
    ('可用现金', '¥80,934'),
]

for item, value in summary:
    p = doc.add_paragraph()
    p.add_run(f'{item}: ').bold = True
    p.add_run(value)

# 五、决策建议
doc.add_heading('五、决策建议', level=1)

decisions = [
    ('克莱机电', '持有', '置信度 30%', '方向不明'),
    ('中国海油', '卖出', '置信度 65%', '多数指标偏空'),
    ('达仁堂', '卖出', '置信度 65%', '多数指标偏空'),
    ('葵花药业', '卖出', '置信度 65%', '多数指标偏空'),
]

table3 = doc.add_table(rows=5, cols=4)
table3.style = 'Table Grid'

header3 = table3.rows[0].cells
header3[0].text = '股票'
header3[1].text = '决策'
header3[2].text = '置信度'
header3[3].text = '原因'

for i, d in enumerate(decisions):
    row = table3.rows[i+1].cells
    for j, cell_data in enumerate(d):
        row[j].text = cell_data

# 六、系统功能清单
doc.add_heading('六、系统功能清单', level=1)

features = [
    '✅ AI因子挖掘 (ai_factor_miner.py)',
    '✅ AI策略生成 (ai_strategy_generator.py)',
    '✅ 财报分析 (financial_analyzer.py)',
    '✅ AkShare数据集成 (akshare_data.py)',
    '✅ 增强风控 (risk_manager.py)',
    '✅ Playwright爬虫 (playwright_crawler.py)',
    '✅ 决策引擎 (decision_engine.py)',
    '✅ 模拟交易 (simulated_trading.py)',
    '✅ 交易机器人 (trading_bot.py)',
]

for feature in features:
    doc.add_paragraph(feature)

# 七、建议
doc.add_heading('七、建议', level=1)

suggestions = [
    '1. 解决AkShare个股K线限制问题（手动导出或换网络）',
    '2. 补充历史数据到3年以上',
    '3. 开启定时任务自动运行',
    '4. 考虑实盘对接',
]

for s in suggestions:
    doc.add_paragraph(s)

# 八、风险提示
doc.add_heading('八、风险提示', level=1)
doc.add_paragraph('⚠️ 本报告仅供参考，不构成投资建议。交易有风险，投资需谨慎。')

# 页脚
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('🦞 量化交易系统 v3.0 - 试运行报告').bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 保存
output_path = os.path.expanduser('~/quant_system/量化交易系统v3.0试运行报告.docx')
doc.save(output_path)
print(f'✅ Word文档已生成: {output_path}')
